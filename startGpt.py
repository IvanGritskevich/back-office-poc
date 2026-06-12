import os
from google import genai
from google.genai import types
from dotenv import load_dotenv
from pathlib import Path
from typing import Optional
import asyncpg
from gpt import InvoiceData
import io
import docx
import openpyxl
from groq import AsyncGroq  # Импортируем асинхронный клиент Groq
import json

async def save_to_pending(sender_email: str, result) -> int:
    """Безопасное сохранение сырых данных ИИ во временную таблицу"""
    
    # Конвертируем Pydantic-модель в JSON строку
    if hasattr(result, "model_dump_json"):
        json_data = result.model_dump_json()
    else:
        json_data = result.json()

    # Открываем СВЕЖЕЕ соединение прямо в момент записи
    # Переменные окружения берутся из твоего key.env, который загружается при старте
    conn = await asyncpg.connect(
        user=os.getenv("DB_USER", "postgres"),
        password=os.getenv("DB_PASSWORD"),
        database=os.getenv("DB_NAME", "exel_group"),
        host=os.getenv("DB_HOST", "postgres_db"),
        port=os.getenv("DB_PORT", "5432")
    )
    
    try:
        # Выполняем запрос внутри блока try
        row = await conn.fetchrow(
            """
            INSERT INTO pending_invoices (sender_email, raw_data)
            VALUES ($1, $2)
            RETURNING id;
            """,
            sender_email, json_data
        )
        return row['id']
        
    finally:
        # Блок finally выполнится ЖЕЛЕЗНО, даже если запрос упадет с ошибкой.
        # Это гарантирует, что соединение не «повиснет» в пуле Postgres
        await conn.close()


# Загружаем переменные окружения
dotenv_path = Path('key.env')
load_dotenv(dotenv_path=dotenv_path)

# Инициализируем клиентов ИИ
client = genai.Client(api_key=os.getenv("API_KEY"))
groq_client = AsyncGroq(api_key=os.getenv("API_KEY_GROG"))


def parse_docx(file_bytes: bytes) -> str:
    """Извлекает весь текст из файла Word (.docx)"""
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)


def parse_xlsx(file_bytes: bytes) -> str:
    """Извлекает все данные из таблиц Excel (.xlsx)"""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    full_text = []
    for sheet in wb.worksheets:
        full_text.append(f"--- Лист: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(row):
                row_text = [str(cell) if cell is not None else "" for cell in row]
                full_text.append(" | ".join(row_text))
    return "\n".join(full_text)


async def save_to_db(result: InvoiceData, manager_id: int):
    conn = await asyncpg.connect(
        user=os.getenv("DB_USER"), 
        password=os.getenv("DB_PASSWORD"),
        database=os.getenv("DB_NAME"), 
        host=os.getenv("DB_HOST")
    )
    
    try:
        # Открываем транзакцию, чтобы всё записалось атомарно
        async with conn.transaction():
            # 1. Проверяем, существует ли уже клиент с такой почтой
            client_id = None
            if result.email:
                client_id = await conn.fetchval("SELECT client_id FROM clients WHERE email = $1", result.email)
                
            # Если клиента нет, создаем новую карточку клиента
            if not client_id:
                address_str = result.address.to_string() if result.address else "Не указан"
                client_id = await conn.fetchval("""
                    INSERT INTO clients (name, address, city, country, postal, email) 
                    VALUES ($1, $2, $3, $4, $5, $6) 
                    RETURNING client_id
                """, result.name, address_str, result.city, result.country, result.postal, result.email)

            # Превращаем "Да"/"Нет" от ИИ в понятный для Postgres BOOLEAN (True/False)
            is_nz_tax = True if result.has_nz_tax_15 == "Да" else False

            # 2. Записываем сам счет в таблицу bills, привязывая его к ID клиента и ID менеджера
            await conn.execute("""
                INSERT INTO bills (client_id, amount, currency, has_nz_tax_15, created_by) 
                VALUES ($1, $2, $3, $4, $5)
            """, client_id, result.amount, result.currency, is_nz_tax, manager_id)
        
    finally:
        await conn.close()

async def extract_invoice_multimedia(text_prompt: str, files_list: list) -> InvoiceData:
    """
    Принимает накопленный текст и список кортежей файлов [(file_bytes, mime_type), ...]
    И отправляет всё ОДНИМ запросом в ИИ.
    """
    contents = []
    
    # 1. Проходимся по всем накопленным файлам и парсим их в зависимости от типа
    for file_bytes, mime_type in files_list:
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_text = parse_docx(file_bytes)
            contents.append(f"\n[Данные из Word документа]:\n{extracted_text}")
        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            extracted_text = parse_xlsx(file_bytes)
            contents.append(f"\n[Данные из Excel таблицы]:\n{extracted_text}")
        elif mime_type:
            # Для PDF и Картинок создаем Part объект (Gemini «увидит» их одновременно)
            file_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            contents.append(file_part)
            
    # 2. Добавляем объединенный текст переписки или подписей
    if text_prompt:
        contents.append(f"\nСопроводительный текст / Подписи: {text_prompt}")

    # Строгая инструкция по валидации данных (одинакова для всех моделей)
    system_instruction = (
        "Ты — эксперт по верификации международных адресов и парсингу данных.\n\n"
        "ПРАВИЛА И КРИТЕРИИ ВАЛИДАЦИИ:\n"
        "1. СТРУКТУРА АДРЕСА: Внутри объекта address разделяй данные строго по полям:\n"
        "   - address.street: Только улица, дом, корпус, квартира (например: 'Hacharimon Street 20, Apartment 6'). НЕ пиши сюда индекс, город и страну.\n"
        "   - address.city: Только город (например: 'Safed').\n"
        "   - address.country: Только страна (например: 'Israel').\n"
        "2. КОРНЕВЫЕ ПОЛЯ: Обязательно продублируй город в корневое поле 'city', а страну — в корневое поле 'country'. Они НЕ должны быть путыми.\n"
        "3. ПОЧТОВЫЙ ИНДЕКС: Найди почтовый индекс (ZIP/Postal code) и запиши его строго в отдельное поле 'postal'. Не удаляй его.\n"
        "4. ФИНАНСЫ: Сумму пиши в 'amount' (float), код валюты — в 'currency' (например, 'USD', 'EUR', 'ILS'). Если данных нет, ставь 0.0 и 'USD'.\n"
        "5. Если страна 'New Zealand' -> 'has_nz_tax_15' = 'Да', иначе 'Нет'.\n\n"
        
        "ПРИМЕР ИДЕАЛЬНОГО РАЗБОРА:\n"
        "Входной текст: 'Anastasia Katsevman, Hahermon St 20, Apt 6, Safed, Israel, 1310401, Nastya Belikova-Kats, Freezingboo@gmail.com'\n"
        "Твой JSON-выход должен быть точно таким:\n"
        "{\n"
        "  'name': 'Anastasia Katsevman',\n"
        "  'address': {\n"
        "    'street': 'Hacharimon Street 20, Apartment 6',\n"
        "    'city': 'Safed',\n"
        "    'country': 'Israel'\n"
        "  },\n"
        "  'city': 'Safed',\n"
        "  'country': 'Israel',\n"
        "  'username': 'Nastya Belikova-Kats',\n"
        "  'email': 'Freezingboo@gmail.com',\n"
        "  'amount': 0.0,\n"
        "  'currency': 'USD',\n"
        "  'postal': '1310401',\n"
        "  'has_nz_tax_15': 'Нет'\n"
        "}"
    )

    # Каскадный вызов ИИ (План А -> Б -> В)
    try:
        print(f"🤖 [План А] Анализируем медиапакет ({len(files_list)} файлов) в Gemini 2.5 Flash...")
        response = await client.aio.models.generate_content(
            model='gemini-2.5-flash', 
            contents=contents, 
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=InvoiceData, 
                system_instruction=system_instruction
            ),
        )
        return InvoiceData.model_validate_json(response.text)
        
    except Exception as e_gemini_25:
        print(f"⚠️ План А сбой: {e_gemini_25}. Пробуем План Б (Gemini 2.0)...")
        try:
            response = await client.aio.models.generate_content(
                model='gemini-2.0-flash', 
                contents=contents, 
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=InvoiceData, 
                    system_instruction=system_instruction
                ),
            )
            return InvoiceData.model_validate_json(response.text)
        except Exception as e_gemini_20:
            print(f"🚨 План Б сбой: {e_gemini_20}. Переключаемся на План В (Groq)...")
            
            # Текстовая сборка для Groq (так как он не видит картинки/PDF напрямую)
            text_context = f"{system_instruction}\n\nСобери данные из текстов в единый JSON:\n"
            for item in contents:
                if isinstance(item, str):
                    text_context += f"\n{item}"
                    
            groq_response = await groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": text_context}],
                response_format={"type": "json_object"}
            )
            return InvoiceData.model_validate_json(groq_response.choices[0].message.content)